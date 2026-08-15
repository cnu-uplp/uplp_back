from datetime import datetime, timedelta, timezone
from io import BytesIO

from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.exc import IntegrityError
from sqlalchemy.orm import Session

from config import settings
from database import get_db
from deps import (
    APPROVAL_APPROVED,
    MEMBERSHIP_STUDENT,
    ROSTER_VISIBLE_MEMBERSHIPS,
    get_current_staff,
    get_current_user,
    get_current_user_optional,
)
from models import SwimApplication, SwimSession, User
from schemas import SwimApplyRequest, SwimCapacityUpdate, SwimSessionCreate

router = APIRouter(prefix="/api/swim", tags=["swim"])

DIVISIONS = ("training", "progress")
DIVISION_LABELS = {"training": "훈련부", "progress": "진도부"}


def _now() -> datetime:
    return datetime.utcnow()


def _parse_iso_utc(value: str) -> datetime:
    """프론트의 ISO 문자열(UTC)을 naive UTC datetime으로 변환."""
    try:
        dt = datetime.fromisoformat(value.replace("Z", "+00:00"))
    except ValueError:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "잘못된 날짜 형식입니다.")
    if dt.tzinfo is not None:
        dt = dt.astimezone(timezone.utc).replace(tzinfo=None)
    return dt


def _session_status(s: SwimSession, now: datetime) -> str:
    if now < s.apply_start_at:
        return "upcoming"  # 오픈 예정 (목록엔 뜨지만 신청 불가)
    if now <= s.apply_end_at:
        return "open"      # 신청 중 (선착순)
    return "closed"        # 마감


def _division_board(s: SwimSession, apps: list[SwimApplication], division: str):
    """부서별 현황 계산. 순번은 저장하지 않고 매번 계산
    → 취소가 생기면 자동으로 뒤 사람이 한 칸씩 당겨진다.

    순서 규칙: (일반 큐 전체, 선착순) → (병합된 후순위, 선착순).
    후순위는 병합돼도 queue="late" 출신 표식을 유지하므로,
    배정이든 예비번호든 일반 큐 인원이 후순위 인원보다 '항상' 앞선다
    (병합 후에 신청한 일반 회원도 후순위보다 앞).
    - assigned    : 정원 내 배정
    - reserve     : 예비번호 (정원 초과분)
    - pending_late: 후순위 대기열 (병합 전 — 배정/예비에 안 들어감)"""
    cap = s.cap_training if division == "training" else s.cap_progress
    d = [a for a in apps if a.division == division]
    key = lambda a: (a.applied_at, a.id)  # noqa: E731
    normal = sorted([a for a in d if a.queue == "normal"], key=key)
    merged_late = sorted([a for a in d if a.queue == "late" and a.merged], key=key)
    pending_late = sorted([a for a in d if a.queue == "late" and not a.merged], key=key)
    ordered = normal + merged_late
    return cap, ordered[:cap], ordered[cap:], pending_late


def _serialize_session(s: SwimSession, apps: list[SwimApplication], now: datetime, user: User | None):
    counts = {}
    my = None
    for div in DIVISIONS:
        cap, assigned, reserve, late = _division_board(s, apps, div)
        counts[div] = {
            "cap": cap,
            "assigned": len(assigned),
            "reserve": len(reserve),
            "pendingLate": len(late),
        }
        if user is not None:
            for i, a in enumerate(assigned):
                if a.user_id == user.id:
                    my = {"division": div, "state": "assigned", "rank": i + 1}
            for i, a in enumerate(reserve):
                if a.user_id == user.id:
                    my = {"division": div, "state": "reserve", "rank": i + 1}
            for i, a in enumerate(late):
                if a.user_id == user.id:
                    my = {"division": div, "state": "pending_late", "rank": i + 1}
    return {
        "id": s.id,
        "meetDate": s.meet_date,
        "meetTime": s.meet_time,
        "location": s.location,
        "capTraining": s.cap_training,
        "capProgress": s.cap_progress,
        "lateQueueEnabled": s.late_queue_enabled,
        "applyStartAt": s.apply_start_at.isoformat() + "Z",
        "applyEndAt": s.apply_end_at.isoformat() + "Z",
        "status": _session_status(s, now),
        "counts": counts,
        "my": my,
    }


@router.get("/sessions")
def list_sessions(
    db: Session = Depends(get_db),
    user: User | None = Depends(get_current_user_optional),
):
    now = _now()
    sessions = db.query(SwimSession).order_by(SwimSession.id.desc()).all()
    out = []
    for s in sessions:
        apps = db.query(SwimApplication).filter(SwimApplication.session_id == s.id).all()
        out.append(_serialize_session(s, apps, now, user))
    return out


@router.post("/sessions")
def create_session(
    payload: SwimSessionCreate,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """정기수영 열기 — 관리자 전용. 신청 시작 전엔 '오픈 예정'으로만 노출된다."""
    start = _parse_iso_utc(payload.applyStartAt)
    end = _parse_iso_utc(payload.applyEndAt)
    if end <= start:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "마감 시각은 시작 시각보다 뒤여야 합니다.")
    if payload.capTraining < 0 or payload.capProgress < 0:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "인원은 0 이상이어야 합니다.")
    if not payload.meetDate or not payload.meetTime or not payload.location.strip():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "모이는 날·시각·위치를 입력해주세요.")

    s = SwimSession(
        meet_date=payload.meetDate,
        meet_time=payload.meetTime,
        location=payload.location.strip(),
        cap_training=payload.capTraining,
        cap_progress=payload.capProgress,
        late_queue_enabled=payload.lateQueueEnabled,
        apply_start_at=start,
        apply_end_at=end,
    )
    db.add(s)
    db.commit()
    db.refresh(s)
    return _serialize_session(s, [], _now(), admin)


@router.post("/sessions/{sid}/apply")
def apply(
    sid: int,
    payload: SwimApplyRequest,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """선착순 신청. 후순위 상태 회원(+제도 적용 세션)은 후순위 대기열로 들어간다."""
    # 재학생만 신청할 수 있다. 졸업생·외부인은 프론트에서도 버튼을 숨기지만,
    # 개발자 도구로 우회해도 여기서 막힌다.
    # (대관 명단에 이름+전화번호가 들어가는데 학생 외에는 연락처를 받지 않기 때문)
    if user.membership != MEMBERSHIP_STUDENT:
        raise HTTPException(
            status.HTTP_403_FORBIDDEN,
            "정기수영 신청은 재학생 부원만 가능합니다.",
        )
    # 임원진 승인 전에는 신청할 수 없다 (둘러보기는 가능).
    if user.approval_status != APPROVAL_APPROVED:
        raise HTTPException(
            status.HTTP_403_FORBIDDEN,
            "임원진 승인 후 신청할 수 있습니다.",
        )

    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    if payload.division not in DIVISIONS:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "부서는 훈련부/진도부 중 하나여야 합니다.")

    now = _now()
    st = _session_status(s, now)
    if st == "upcoming":
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "아직 신청 시작 전입니다.")
    if st == "closed":
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "신청이 마감되었습니다.")

    existing = (
        db.query(SwimApplication)
        .filter(SwimApplication.session_id == sid, SwimApplication.user_id == user.id)
        .first()
    )
    if existing:
        raise HTTPException(status.HTTP_409_CONFLICT, "이미 신청했습니다. (취소 후 다시 신청 가능)")

    queue = "late" if (s.late_queue_enabled and user.is_deprioritized) else "normal"
    db.add(
        SwimApplication(
            session_id=sid, user_id=user.id, division=payload.division, queue=queue, applied_at=now
        )
    )
    # 위의 existing 조회와 이 INSERT 사이는 원자적이지 않다.
    # 더블클릭·재시도로 요청이 둘 동시에 들어오면 양쪽 다 "신청 없음"으로 판단해
    # 둘 다 INSERT 하고, uq_swim_app_session_user 제약에 걸려 뒤엣것이 터진다.
    # 그대로 두면 500이 나가서 "실패한 줄 알고 또 누르는" 사고가 난다 — 409로 돌려준다.
    try:
        db.commit()
    except IntegrityError:
        db.rollback()
        raise HTTPException(
            status.HTTP_409_CONFLICT, "이미 신청했습니다. (취소 후 다시 신청 가능)"
        )

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    return _serialize_session(s, apps, now, user)


@router.delete("/sessions/{sid}/apply")
def cancel(
    sid: int,
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """취소 — 즉시 반영. 순번은 매번 계산하므로 뒷사람이 자동으로 당겨진다."""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    now = _now()
    if _session_status(s, now) == "closed":
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "마감된 정기수영은 취소할 수 없습니다.")

    row = (
        db.query(SwimApplication)
        .filter(SwimApplication.session_id == sid, SwimApplication.user_id == user.id)
        .first()
    )
    if row is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "신청 내역이 없습니다.")
    db.delete(row)
    db.commit()

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    return _serialize_session(s, apps, now, user)


@router.post("/sessions/{sid}/merge")
def merge(
    sid: int,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """병합 — 관리자 전용. 후순위 대기열을 FIFO 그대로 빈자리·예비번호 계산에 합류시킨다.
    출신 표식(queue="late")은 유지되므로, 배정·예비 모두에서
    일반 큐 인원이 후순위 인원보다 항상 앞선다."""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")

    late_apps = (
        db.query(SwimApplication)
        .filter(
            SwimApplication.session_id == sid,
            SwimApplication.queue == "late",
            SwimApplication.merged == False,  # noqa: E712
        )
        .order_by(SwimApplication.applied_at, SwimApplication.id)
        .all()
    )
    for a in late_apps:
        a.merged = True
    db.commit()

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    result = _serialize_session(s, apps, _now(), admin)
    result["mergedCount"] = len(late_apps)
    return result


@router.patch("/sessions/{sid}")
def update_session(
    sid: int,
    payload: SwimSessionCreate,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """정기수영 수정 — 관리자 전용, '신청이 열리기 전(오픈 예정)'에만 가능.
    신청이 시작된 뒤에는 선착순 공정성 때문에 조건을 바꿀 수 없다."""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    if _session_status(s, _now()) != "upcoming":
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "신청이 시작된 뒤에는 수정할 수 없습니다.",
        )

    start = _parse_iso_utc(payload.applyStartAt)
    end = _parse_iso_utc(payload.applyEndAt)
    if end <= start:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "마감 시각은 시작 시각보다 뒤여야 합니다.")
    if payload.capTraining < 0 or payload.capProgress < 0:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "인원은 0 이상이어야 합니다.")
    if not payload.meetDate or not payload.meetTime or not payload.location.strip():
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "모이는 날·시각·위치를 입력해주세요.")

    s.meet_date = payload.meetDate
    s.meet_time = payload.meetTime
    s.location = payload.location.strip()
    s.cap_training = payload.capTraining
    s.cap_progress = payload.capProgress
    s.late_queue_enabled = payload.lateQueueEnabled
    s.apply_start_at = start
    s.apply_end_at = end
    db.commit()
    db.refresh(s)

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    return _serialize_session(s, apps, _now(), admin)


@router.patch("/sessions/{sid}/capacity")
def update_capacity(
    sid: int,
    payload: SwimCapacityUpdate,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """인원(정원)만 수정 — 관리자 전용. 신청을 받는 중에도 조절할 수 있다.

    배정/예비는 저장하지 않고 매번 순서대로 계산하므로:
    - 정원을 줄이면 뒤쪽 인원이 자동으로 예비(명단 대기)로 밀린다
    - 정원을 늘리면 예비 앞쪽 인원이 자동으로 배정으로 올라온다
    (시각·위치 등 다른 조건은 선착순 공정성 때문에 신청 중 변경 불가)"""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    if _session_status(s, _now()) == "closed":
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "마감된 뒤에는 인원을 수정할 수 없습니다.")
    if payload.capTraining < 0 or payload.capProgress < 0:
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "인원은 0 이상이어야 합니다.")

    s.cap_training = payload.capTraining
    s.cap_progress = payload.capProgress
    db.commit()
    db.refresh(s)

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    return _serialize_session(s, apps, _now(), admin)


@router.post("/sessions/{sid}/close")
def close_session(
    sid: int,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """즉시 마감 — 관리자 전용. 마감 시각을 지금으로 당겨 신청·취소를 막고
    명단 다운로드를 열어준다."""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    now = _now()
    if _session_status(s, now) == "closed":
        raise HTTPException(status.HTTP_400_BAD_REQUEST, "이미 마감되었습니다.")
    s.apply_end_at = now
    # 아직 시작 전이었다면 시작 시각도 당겨 '마감' 상태가 되도록 정리
    if s.apply_start_at > now:
        s.apply_start_at = now
    db.commit()

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    return _serialize_session(s, apps, _now(), admin)


@router.delete("/sessions/{sid}")
def delete_session(
    sid: int,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """정기수영 삭제 — 관리자 전용. 잘못 만든 회차를 신청 내역과 함께 제거한다."""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    db.query(SwimApplication).filter(SwimApplication.session_id == sid).delete()
    db.delete(s)
    db.commit()
    return {"ok": True, "deleted": sid}


@router.get("/sessions/{sid}/roster")
def roster(
    sid: int,
    db: Session = Depends(get_db),
    user: User | None = Depends(get_current_user_optional),
):
    """명단 대시보드 — 누구나 '몇 명 신청했는지'는 볼 수 있다.

    다만 실명은 개인정보이므로 재학생·졸업생(로그인 상태)에게만 보인다.
    외부인과 비로그인에게는 "***" 로 마스킹한다.
    (연락처·학과는 어느 경우에도 내보내지 않는다 — 관리자 docx 전용)"""
    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")

    can_see_names = (
        user is not None
        and user.membership in ROSTER_VISIBLE_MEMBERSHIPS
        and user.approval_status == APPROVAL_APPROVED  # 승인 전에는 외부인과 같이 마스킹
    )

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    user_ids = {a.user_id for a in apps}
    users = (
        {u.id: u for u in db.query(User).filter(User.id.in_(user_ids)).all()} if user_ids else {}
    )

    def name_of(a: SwimApplication) -> str:
        if not can_see_names:
            return "***"
        u = users.get(a.user_id)
        if u is None:
            return f"회원{a.user_id}"
        # 실명(name) 우선 — 카카오 닉네임은 대관 명단에 쓸 수 없다
        return u.name or u.nickname or f"회원{u.id}"

    divisions = {}
    for div in DIVISIONS:
        cap, assigned, reserve, pending = _division_board(s, apps, div)
        divisions[div] = {
            "cap": cap,
            "assigned": [{"rank": i + 1, "name": name_of(a)} for i, a in enumerate(assigned)],
            "reserve": [{"rank": i + 1, "name": name_of(a)} for i, a in enumerate(reserve)],
            "pendingLate": [{"rank": i + 1, "name": name_of(a)} for i, a in enumerate(pending)],
        }
    return {"sessionId": sid, "divisions": divisions}


def _kst(dt: datetime) -> str:
    return (dt + timedelta(hours=9)).strftime("%Y-%m-%d %H:%M")


@router.get("/sessions/{sid}/roster.docx")
def roster_docx(
    sid: int,
    admin: User = Depends(get_current_staff),
    db: Session = Depends(get_db),
):
    """명단 docx 다운로드 — 관리자 전용, 신청 '마감 후'에만 가능."""
    # 무거운 import는 사용 시점에
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Mm, Pt, RGBColor

    s = db.get(SwimSession, sid)
    if s is None:
        raise HTTPException(status.HTTP_404_NOT_FOUND, "정기수영을 찾을 수 없습니다.")
    if _session_status(s, _now()) != "closed":
        raise HTTPException(
            status.HTTP_400_BAD_REQUEST,
            "명단은 신청 마감 후에 다운로드할 수 있습니다.",
        )

    apps = db.query(SwimApplication).filter(SwimApplication.session_id == sid).all()
    user_ids = {a.user_id for a in apps}
    users = (
        {u.id: u for u in db.query(User).filter(User.id.in_(user_ids)).all()} if user_ids else {}
    )

    def display(a: SwimApplication) -> tuple[str, str, str]:
        u = users.get(a.user_id)
        if u is None:
            return (f"회원{a.user_id}", "-", "-")
        name = u.name or u.nickname or f"회원{u.id}"
        phone = u.phone_number or "-"
        dept = " ".join(x for x in (u.college, u.department) if x) or "-"
        return (name, phone, dept)

    # 훈련부·진도부 '배정' 인원만 합쳐 한 표로 (예비·후순위 대기는 명단에 미포함)
    roster: list[SwimApplication] = []
    for div in DIVISIONS:
        _cap, assigned, _reserve, _late = _division_board(s, apps, div)
        roster.extend(assigned)

    # ── 스포렉스 레인대관 신청서 양식 재현 ──────────────────────────────
    # 아래 좌표·크기는 전부 원본 PDF(pdfplumber 실측)에서 뽑은 값이다.
    # 단위는 mm, 페이지는 A4(210×297).
    #
    # 원본은 한글(HWP) 문서라 줄간격이 글자크기의 160%인데,
    # 워드의 맑은 고딕 기본 줄간격은 174%라 그냥 두면 아래로 계속 밀린다.
    # → 모든 문단·셀에 "고정(EXACTLY) 줄간격 = 글자크기 × 1.6" 을 준다.
    # 이때 글리프 윗변은  line_top + L − 1.2014×F  에 온다(실측 검증됨).
    MM_PT = 1 / 0.352778  # mm → pt

    def pt(mm_val):
        return Pt(mm_val * MM_PT)

    doc = Document()

    # 1페이지 섹션 — 원본 실측: 위 7.07mm, 좌우 10mm, 아래 12mm, 쪽번호 없음
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Mm(10.0)
    section.top_margin = Mm(7.07)
    # 1페이지는 본문이 283.8mm 까지 차므로 아래 여백을 조금 더 줄여 여유를 둔다
    section.bottom_margin = Mm(10.5)

    # 기본 글꼴 (한글 폰트는 eastAsia 속성까지 지정해야 적용된다)
    normal = doc.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal.font.size = Pt(10)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    def no_autospace(p):
        """한글과 숫자·영문 사이에 워드가 넣는 자동 간격을 끈다 (원본 hwp처럼 붙여 쓰기)."""
        pPr = p._p.get_or_add_pPr()
        for tag in ("w:autoSpaceDE", "w:autoSpaceDN"):
            el = pPr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                pPr.append(el)
            el.set(qn("w:val"), "0")

    def shade(cell, hex_color):
        """셀 배경색 (원본 양식의 색을 그대로 재현)."""
        el = OxmlElement("w:shd")
        el.set(qn("w:val"), "clear")
        el.set(qn("w:color"), "auto")
        el.set(qn("w:fill"), hex_color)
        cell._tc.get_or_add_tcPr().append(el)

    # 원본 표의 셀 좌우 안여백 실측: 1.9mm (= 108 dxa)
    CELL_PAD_DXA = 108
    CELL_PAD_MM = 1.9

    def set_cell_font(cell, size=10, bold=False, align=None, color=None,
                      line_pt=None, vcenter=True):
        """셀 글자 — 고정 줄간격 + 세로 가운데. 원본 y좌표와 일치하도록."""
        if vcenter:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for p in cell.paragraphs:
            no_autospace(p)
            if align is not None:
                p.alignment = align
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = Pt(line_pt if line_pt is not None else size * 1.6)
            for run in p.runs:
                run.font.size = Pt(size)
                run.font.name = "맑은 고딕"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
                run.bold = bold
                if color is not None:
                    run.font.color.rgb = color

    def set_rows(table, heights_mm):
        """행 높이를 원본 실측값으로 '정확히' 고정한다.

        cantSplit 을 함께 켜서 행이 페이지 경계에서 반으로 잘리지 않게 한다.
        (trPr 자식 순서상 cantSplit 은 trHeight 보다 앞에 와야 한다.)"""
        for row, h in zip(table.rows, heights_mm):
            row.height = Mm(h)
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            trPr = row._tr.get_or_add_trPr()
            if trPr.find(qn("w:cantSplit")) is None:
                trPr.insert(0, OxmlElement("w:cantSplit"))

    # 원본 실측 선 굵기 (w:sz 단위 = 1/8 pt)
    #   바깥 1.44pt → 12,  명단표 위·아래 1.2pt → 10,
    #   명단표 헤더 아래 0.48pt → 4,  안쪽 칸선 0.24pt → 2
    SZ_OUTER, SZ_ROSTER, SZ_HEADER, SZ_INNER = 12, 10, 4, 2

    def cell_border(cell, **sides):
        """셀 단위 테두리. sides 예) top=12, bottom=12"""
        tcPr = cell._tc.get_or_add_tcPr()
        el = tcPr.find(qn("w:tcBorders"))
        if el is None:
            el = OxmlElement("w:tcBorders")
            tcPr.append(el)
        for side, sz in sides.items():
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "single" if sz else "none")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            el.append(b)

    def style_table(table, widths_mm, left_mm, borders=True,
                    outer_sz=None, inner_sz=None):
        """표 폭·위치·테두리 고정.

        left_mm 은 '원본 PDF에서 잰 표 왼쪽 테두리의 x좌표'다.
        워드는 표 왼쪽 테두리를 (본문여백 + tblInd − 셀좌여백) 에 놓으므로
        그 관계를 역산해서 tblInd 를 구한다(실측으로 검증한 규칙).
        """
        page_left = section.left_margin.mm
        indent_mm = left_mm - page_left + CELL_PAD_MM

        widths = [Mm(w) for w in widths_mm]
        table.autofit = False
        for i, w in enumerate(widths):
            if i < len(table.columns):
                table.columns[i].width = w
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(widths):
                    cell.width = widths[i]
        pr = table._tbl.tblPr
        for tag, attrs in (
            ("w:tblW", {"w:type": "dxa", "w:w": str(sum(w.twips for w in widths))}),
            ("w:tblInd", {"w:type": "dxa", "w:w": str(int(round(indent_mm * 56.6929)))}),
            ("w:tblLayout", {"w:type": "fixed"}),
        ):
            el = pr.find(qn(tag))
            if el is None:
                el = OxmlElement(tag)
                pr.append(el)
            for k, v in attrs.items():
                el.set(qn(k), v)
        # 셀 상하 여백 0 — 행 높이가 불필요하게 커지는 것을 막는다
        cell_mar = OxmlElement("w:tblCellMar")
        for side, val in (("top", "0"), ("bottom", "0"),
                          ("left", str(CELL_PAD_DXA)), ("right", str(CELL_PAD_DXA))):
            m = OxmlElement(f"w:{side}")
            m.set(qn("w:w"), val)
            m.set(qn("w:type"), "dxa")
            cell_mar.append(m)
        pr.append(cell_mar)

        # 원본은 바깥 테두리만 굵고(1.44pt) 안쪽 칸선은 가늘다(0.24pt)
        o_sz = outer_sz if outer_sz is not None else SZ_OUTER
        i_sz = inner_sz if inner_sz is not None else SZ_INNER

        def edge_size(edge):
            if not borders:
                return 0
            if edge in ("insideH", "insideV"):
                return i_sz
            return o_sz[edge] if isinstance(o_sz, dict) else o_sz

        tbl_borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            sz = edge_size(edge)
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single" if sz else "none")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000" if sz else "auto")
            tbl_borders.append(el)
        pr.append(tbl_borders)

        # OOXML은 자식 요소 '순서'까지 스키마를 따라야 한다.
        # 순서가 틀리면 워드가 "파일을 열 때 오류" 를 내며 거부한다.
        TBLPR_ORDER = [
            "tblStyle", "tblpPr", "tblOverlap", "bidiVisual", "tblStyleRowBandSize",
            "tblStyleColBandSize", "tblW", "jc", "tblCellSpacing", "tblInd", "tblBorders",
            "shd", "tblLayout", "tblCellMar", "tblLook", "tblCaption", "tblDescription",
            "tblPrChange",
        ]
        children = list(pr)
        children.sort(key=lambda c: TBLPR_ORDER.index(c.tag.split("}")[1])
                      if c.tag.split("}")[1] in TBLPR_ORDER else len(TBLPR_ORDER))
        for c in children:
            pr.append(c)  # append는 기존 노드를 뒤로 옮긴다 → 정렬 순서대로 재배치
        return table

    def para(text="", size=10, bold=False, align=None, after_mm=0.0,
             before_mm=0.0, indent_mm=0.0, right_mm=0.0,
             prefix=None, prefix_size=None, line_pt=None):
        """원본 PDF 실측값(폰트 크기·들여쓰기·간격)을 그대로 재현한다.

        before_mm / after_mm 은 원본에서 잰 '문단 사이 거리'를 mm 로 넣는다.
        line_pt 를 주지 않으면 고정 줄간격 = 글자크기 × 1.6 (한글 문서 기본값).
        prefix: ①②③ 처럼 기호만 크기가 다른 경우 사용."""
        p = doc.add_paragraph()
        no_autospace(p)
        pf = p.paragraph_format
        pf.space_before = pt(before_mm)
        pf.space_after = pt(after_mm)
        pf.line_spacing = Pt(line_pt if line_pt is not None
                             else max(size, prefix_size or 0) * 1.6)
        if indent_mm:
            pf.left_indent = Mm(indent_mm)
        if right_mm:
            pf.right_indent = Mm(right_mm)
        if align is not None:
            p.alignment = align

        def _mk(t, sz):
            r = p.add_run(t)
            r.font.size = Pt(sz)
            r.font.name = "맑은 고딕"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
            r.bold = bold
            return r

        if prefix is not None:
            _mk(prefix, prefix_size or size)
        _mk(text, size)
        return p

    CENTER = WD_ALIGN_PARAGRAPH.CENTER
    RIGHT = WD_ALIGN_PARAGRAPH.RIGHT

    # ── 1페이지: 신청서 ─────────────────────────────────────────────
    # (원본 y좌표) 제목 9.6 / 1. 26.1 / 1-1. 32.3 / 2. 44.7 / 3. 57.1
    para("CNU SPOREX Swimming 레인대관 신청서", size=18, bold=True, align=CENTER,
         after_mm=7.32)

    para("1. 충남대학교 CNU SPOREX Swimming(이하 센터) 레인사용은 1시간 사용을 기준으로 합니다.",
         size=11, bold=True, after_mm=0)
    para("1-1. 이외의 추가 시간 사용은 별도 협의해 주시기 바랍니다.",
         size=11, bold=True, indent_mm=5.8, after_mm=6.19)
    para("2. 당일 취소피해(노쇼)를 최소화 하고자 노쇼발생 시 향후 레인대관 신청을 허가하지 않겠습니다.",
         size=11, bold=True, after_mm=6.19)
    para("3. 센터 사용료는 아래와 같습니다.", size=11, bold=True, after_mm=1.44)

    # 요금표 — 원본: 좌 11.1, 폭 70.9/54.9/63.0, 행 7.03/9.82, y 63.2~80.05
    fee = doc.add_table(rows=2, cols=3)
    fee.cell(0, 0).text, fee.cell(0, 1).text, fee.cell(0, 2).text = "구     분", "금액", "비고"
    fee.cell(1, 0).text, fee.cell(1, 1).text = "1시간 이용", "120,000원"
    fee.cell(1, 2).text = "⚫ 일일 입장료는 별도"
    set_rows(fee, [7.03, 9.82])
    style_table(fee, [70.9, 54.9, 63.0], left_mm=11.1)
    for r in range(2):
        for c in range(3):
            # 원본에서 '비고' 값 칸만 9pt, 나머지는 10pt. 전부 가운데 정렬.
            set_cell_font(fee.cell(r, c), size=9 if (r, c) == (1, 2) else 10,
                          bold=(r == 0), align=CENTER)

    # 원본 y: 4. 87.5 / ① 94.5 / ② 100.8 / 업장에 107.2 / ③ 112.0 / 5. 123.0
    para("4. 대관준수 사항", size=12, bold=True, before_mm=5.76, after_mm=0.37)
    para('“을”은 “갑”의 “센터 대관규약”에 명시된 제반 조항 및 계약을 준수한다.',
         size=9, bold=True, indent_mm=1.9, prefix="① ", prefix_size=11, after_mm=0.09)
    para('“을”은 대관 중 발생한 시설들의 파괴 또는 훼손 등으로 인한 “갑”의 손해에 대해 즉시 배상을 하여야 한다.',
         size=9, indent_mm=1.9, prefix="② ", prefix_size=11, after_mm=0.47)
    # 이 줄만 원본 줄간격이 유난히 좁다(다음 줄까지 4.8mm) → 줄간격을 직접 지정
    para("업장에 손실을 끼쳤을 경우 민·형사 법적인 모든 책임을 묻는다.",
         size=9, indent_mm=7.9, line_pt=12.8, after_mm=0)
    para("대관 시간은 1시간 기준이며, 희망요일 및 레인사용 시간 추가는 미리 협의한다.",
         size=9, bold=True, indent_mm=1.9, prefix="③ ", prefix_size=11, after_mm=4.79)

    para("5. 대관 이용수칙", size=11, bold=True, indent_mm=1.6, after_mm=5.84)
    rules = [
        "수영장 및 휴게실 이용수칙",
        "1. 다음 사용자를 위해 이용 시간을 준수해 주시기 바랍니다.",
        "2. 미취학 아동 및 키제한 110cm이하시 부모님 필히 동반입장 해주시기 바랍니다. ( 수심 – 2개 레인:1.6m, 6개 레인 1.2m )",
        "3. 사용 중 발생한 안전사고 및 상해는 사용자가 책임을 집니다.",
        "4. 수영장의 특성상 침, 가래 등을 뱉는 행위는 절대 금지합니다.",
        "5. 수영장 내 물을 제외한 음식물 및 음료는 반입을 금지합니다.",
        "6. 퇴실 전 가져오신 용품 및 쓰레기는 정리 부탁드립니다.",
        "7. 안전요원 지시 불이행 시 조기 퇴장 될 수 있습니다.",
        "8. 수영장 내 애완동물 출입을 금지합니다.",
        "9. 사용 중 발생한 분실물은 사용자가 책임을 집니다.",
        "10. 상기 위반에 따른 부상 및 시설물 훼손 시 보수 비용을 해당 고객이 책임집니다.",
        "11. 단체강습, 특강 등 이용 회원들에게 불편함 및 위화감이 느껴질 수 있는 행위는 금지한다.",
        "12. 스타트대 사용 전 반드시 안전 관리자의 관리 감독하에 이용한다.",
        "13. 어떠한 형태로의 촬영은 불허하고 적발 시 즉시 퇴장조치 합니다.",
    ]
    # 이용수칙표 — 원본: 좌 11.1, 폭 183.9, y 133.5~225.65.
    # 행 높이가 들쭉날쭉한 것까지 원본 실측 그대로 옮긴다.
    RULE_ROW_MM = [6.95, 6.18, 7.20, 7.20, 6.10, 6.18, 7.20,
                   7.20, 6.18, 7.11, 6.18, 6.18, 6.18, 6.18]
    rt = doc.add_table(rows=len(rules), cols=1)
    for i, line in enumerate(rules):
        rt.cell(i, 0).text = line
        if i == 0:  # 원본: 검은 배경 + 흰 글씨
            shade(rt.cell(i, 0), "000000")
            set_cell_font(rt.cell(i, 0), size=11, bold=True, align=CENTER,
                          color=RGBColor(0xFF, 0xFF, 0xFF))
        else:
            set_cell_font(rt.cell(i, 0), size=9)
    set_rows(rt, RULE_ROW_MM)
    style_table(rt, [183.9], left_mm=11.1)

    # 6. 레인 대관 신청 — 날짜·인원은 세션 값으로 자동 기입 (원본 y 232.1)
    para("6. 레인 대관 신청", size=10, bold=True, before_mm=5.04, after_mm=1.46)
    try:
        _y, _m, _d = s.meet_date.split("-")
        _hh = s.meet_time.split(":")[0]
        use_date = f"{int(_y)}년 {int(_m)}월 {int(_d)}일 {int(_hh)}시"
    except (ValueError, AttributeError, IndexError):
        use_date = f"{s.meet_date} {s.meet_time}"

    # 신청정보표 — 원본: 좌 11.1, 행 6.94/9.82, y 237.8~254.56
    apply_tbl = doc.add_table(rows=2, cols=5)
    heads = ["구     분", "사용 희망 날짜", "이용 시간", "이용 인원 (15명제한)", "예약자 연락처"]
    vals = ["", use_date, settings.rental_hours, f"{len(roster)}명", settings.club_contact]
    for c, (h, v) in enumerate(zip(heads, vals)):
        apply_tbl.cell(0, c).text = h
        apply_tbl.cell(1, c).text = v
        set_cell_font(apply_tbl.cell(0, c), size=10, bold=True, align=CENTER)
        # 원본은 '사용 희망 날짜'만 10pt, 나머지 값은 9pt
        set_cell_font(apply_tbl.cell(1, c), size=10 if c == 1 else 9, align=CENTER)
    set_rows(apply_tbl, [6.94, 9.82])
    style_table(apply_tbl, [20.4, 52.2, 28.3, 38.5, 44.5], left_mm=11.1)

    # 표와 표 사이를 띄우는 빈 문단 (원본: 254.56 → 266.0)
    para(size=1, line_pt=2, after_mm=10.73)

    # 서명표 — 원본: 좌 125.8, 행 9.06/8.30, y 266.0~283.36 (1페이지 끝)
    sign1 = doc.add_table(rows=2, cols=2)
    sign1.cell(0, 0).text, sign1.cell(0, 1).text = "회원/단체명", settings.club_name
    sign1.cell(1, 0).text, sign1.cell(1, 1).text = "서명(인)", settings.club_signer
    set_rows(sign1, [9.06, 8.30])
    style_table(sign1, [31.0, 38.2], left_mm=125.8)
    for r in range(2):
        set_cell_font(sign1.cell(r, 0), size=10, bold=True, align=CENTER)
        set_cell_font(sign1.cell(r, 1), size=10, align=CENTER)

    # ── 2~3페이지: 이용자 명단 (별도 섹션) ─────────────────────────────
    # 원본은 2페이지부터 위 여백이 10.4mm 이고, 쪽번호도 2페이지부터 붙는다.
    sec2 = doc.add_section()
    # add_section() 은 '섹션 구분 문단'을 1페이지 끝에 하나 남긴다.
    # 그대로 두면 보이지 않는 6mm 짜리 빈 줄이 되어 서명표를 2페이지로 밀어낸다.
    for _p in doc.paragraphs:
        _pPr = _p._p.find(qn("w:pPr"))
        if _pPr is not None and _pPr.find(qn("w:sectPr")) is not None:
            _p.paragraph_format.space_before = Pt(0)
            _p.paragraph_format.space_after = Pt(0)
            _p.paragraph_format.line_spacing = Pt(1)
    sec2.page_width = Cm(21.0)
    sec2.page_height = Cm(29.7)
    sec2.left_margin = sec2.right_margin = Mm(10.0)
    sec2.top_margin = Mm(10.4)
    sec2.bottom_margin = Mm(12.0)
    sec2.footer_distance = Mm(4.26)

    # 쪽번호 "- N -" — 1페이지에는 없고 2페이지부터 (원본 y 288.5, 10pt)
    sec2.footer.is_linked_to_previous = False
    footer_p = sec2.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    footer_p.paragraph_format.line_spacing = Pt(16)
    footer_p.add_run("- ")
    # PAGE 필드: fldChar/instrText는 반드시 각각 run(w:r) 안에 들어가야 한다
    r_begin = footer_p.add_run()
    _b = OxmlElement("w:fldChar"); _b.set(qn("w:fldCharType"), "begin")
    r_begin._r.append(_b)
    r_instr = footer_p.add_run()
    _i = OxmlElement("w:instrText"); _i.set(qn("xml:space"), "preserve"); _i.text = "PAGE"
    r_instr._r.append(_i)
    r_end = footer_p.add_run()
    _e = OxmlElement("w:fldChar"); _e.set(qn("w:fldCharType"), "end")
    r_end._r.append(_e)
    footer_p.add_run(" -")
    for run in footer_p.runs:
        run.font.size = Pt(10)
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # 제목 블록 — 원본: 좌 20.9, 폭 168.1, 주황 바 1.5 / 제목 16.4 / 빨강 바 1.4
    title_tbl = doc.add_table(rows=3, cols=1)
    title_tbl.cell(1, 0).text = "충남대학교 CNU SPOREX Swimming\n레인 대관 이용자 명단"
    shade(title_tbl.cell(0, 0), "EF856E")  # 원본 상단 장식 바(주황)
    shade(title_tbl.cell(2, 0), "E62B33")  # 원본 하단 장식 바(빨강)
    set_rows(title_tbl, [1.5, 16.4, 1.4])
    style_table(title_tbl, [168.1], left_mm=20.9, borders=False)
    # 두 줄이 행을 꽉 채우므로 세로 가운데가 아니라 위 기준으로 둔다
    set_cell_font(title_tbl.cell(1, 0), size=20, bold=True, align=CENTER,
                  line_pt=24.1, vcenter=False)
    para(size=1, line_pt=2, after_mm=2.39)

    # "Ⅰ 레인 대관 참석인원" — 원본: 좌 19.7, 청록 10.4 + 여백 1.9 + 연청록 157.8
    sec_tbl = doc.add_table(rows=1, cols=3)
    sec_tbl.cell(0, 0).text, sec_tbl.cell(0, 1).text = "Ⅰ", ""
    sec_tbl.cell(0, 2).text = "레인 대관 참석인원"
    shade(sec_tbl.cell(0, 0), "0098A0")  # 원본: 청록 배경 + 흰 글씨
    shade(sec_tbl.cell(0, 2), "E7F4F6")  # 원본: 연한 청록
    set_rows(sec_tbl, [9.4])
    style_table(sec_tbl, [10.4, 1.9, 157.8], left_mm=19.7, borders=False)
    # 원본: 연청록 칸 위·아래에만 굵은 가로줄 (x 32.1~189.8, 1.44pt)
    cell_border(sec_tbl.cell(0, 2), top=SZ_OUTER, bottom=SZ_OUTER)
    set_cell_font(sec_tbl.cell(0, 0), size=18, bold=True, align=CENTER,
                  line_pt=21.5, color=RGBColor(0xFF, 0xFF, 0xFF))
    set_cell_font(sec_tbl.cell(0, 2), size=18, bold=True, line_pt=21.5)
    para(size=1, line_pt=2, after_mm=2.29)

    # 명단 표 — 원본: 좌 19.7, 헤더 18.4mm, 본문 7.53mm × 38칸, 글자 12pt
    slots = max(38, len(roster))
    roster_tbl = doc.add_table(rows=slots + 1, cols=4)
    for c, h in enumerate(("연번", "참석자 명단", "전화 번호", "참석 확인")):
        roster_tbl.cell(0, c).text = h
        shade(roster_tbl.cell(0, c), "E5E5FF")  # 원본: 연보라 헤더
        set_cell_font(roster_tbl.cell(0, c), size=12, bold=True, align=CENTER)
    for i in range(slots):
        name, phone, _dept = display(roster[i]) if i < len(roster) else ("", "", "")
        row = roster_tbl.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = name
        row.cells[2].text = phone
        row.cells[3].text = ""  # 현장 서명란
        for c in range(4):
            set_cell_font(row.cells[c], size=12, align=CENTER)
    set_rows(roster_tbl, [18.4] + [7.53] * slots)
    # 원본 명단표: 세로선은 '안쪽 3개'만 있고 좌우 바깥은 아예 선이 없다(투명).
    # 가로는 맨 위 1.2pt, 칸선 0.24pt, 헤더 아래만 0.48pt.
    style_table(roster_tbl, [13.7, 55.6, 69.5, 31.4], left_mm=19.7,
                outer_sz={"top": SZ_ROSTER, "bottom": SZ_ROSTER,
                          "left": 0, "right": 0})
    for c in range(4):
        cell_border(roster_tbl.cell(0, c), bottom=SZ_HEADER)

    # ── 3페이지 꼬리말 — 원본 y: 1. 85.8 / 2. 92.0 / 계약일자 106.8 / 표 116.8
    para("1. 개인 신상정보는 대관 신청한 날짜 일주일 후 파기됩니다.",
         size=11, before_mm=6.08, after_mm=0)
    para("2. 위의 레인 대관 수칙을 모두 확인하였으며 사고 발생 시 이의를 제기하지 않겠습니다.",
         size=11, after_mm=8.73)

    today = _now() + timedelta(hours=9)  # KST
    para(f"계약일자 : {today.year}년 {today.month}월 {today.day}일",
         size=10, align=RIGHT, right_mm=18.9, after_mm=5.76)

    sign2 = doc.add_table(rows=2, cols=2)
    sign2.cell(0, 0).text, sign2.cell(0, 1).text = "회원/단체명", settings.club_name
    sign2.cell(1, 0).text, sign2.cell(1, 1).text = "서명(인)", settings.club_signer
    set_rows(sign2, [9.99, 9.40])
    style_table(sign2, [26.1, 32.1], left_mm=120.8)
    for r in range(2):
        set_cell_font(sign2.cell(r, 0), size=10, bold=True, align=CENTER)
        set_cell_font(sign2.cell(r, 1), size=10, align=CENTER)

    # OOXML은 자식 요소 '순서'까지 스키마를 따라야 하고, 어기면 워드가
    # "파일을 열 때 오류" 를 내며 아예 열지 않는다. 저장 직전에 한 번 정렬한다.
    SCHEMA_ORDER = {
        "tblPr": ["tblStyle", "tblpPr", "tblOverlap", "bidiVisual",
                  "tblStyleRowBandSize", "tblStyleColBandSize", "tblW", "jc",
                  "tblCellSpacing", "tblInd", "tblBorders", "shd", "tblLayout",
                  "tblCellMar", "tblLook", "tblCaption", "tblDescription", "tblPrChange"],
        "tcPr": ["cnfStyle", "tcW", "gridSpan", "hMerge", "vMerge", "tcBorders",
                 "shd", "noWrap", "tcMar", "textDirection", "tcFitText", "vAlign",
                 "hideMark", "tcPrChange"],
        "trPr": ["cnfStyle", "divId", "gridBefore", "gridAfter", "wBefore", "wAfter",
                 "cantSplit", "trHeight", "tblHeader", "tblCellSpacing", "jc",
                 "hidden", "ins", "del", "trPrChange"],
    }
    for tag, order in SCHEMA_ORDER.items():
        for el in doc.element.body.iter(qn(f"w:{tag}")):
            for child in sorted(
                list(el),
                key=lambda c: order.index(c.tag.split("}")[1])
                if c.tag.split("}")[1] in order else len(order),
            ):
                el.append(child)  # append 는 기존 노드를 뒤로 옮긴다 → 정렬됨

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="swim_session_{sid}_roster.docx"'
        },
    )
